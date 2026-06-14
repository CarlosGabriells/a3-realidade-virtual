from copy import deepcopy
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path("/Users/carlosgabriel/Downloads/Cópia de Modelo-de-Projeto-de-Pesquisa ABNT.docx")
OUTPUT = ROOT / "report" / "Relatorio_Projeto_Acessibilidade_VR_ABNT.docx"
ASSETS = ROOT / "report" / "assets"

TITLE = "SISTEMA IMERSIVO DE PRÉ-AUDITORIA DE ACESSIBILIDADE EM AMBIENTES CONSTRUÍDOS"
SUBTITLE = "Aplicação de computação gráfica e realidade virtual baseada em critérios selecionados da ABNT NBR 9050"
ALPHA_NUM_ID = None


def set_cell_shading(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), color)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths_cm):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for index, width in enumerate(widths_cm):
            row.cells[index].width = Cm(width)
            tc_pr = row.cells[index]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(Cm(width).twips)))
            tc_w.set(qn("w:type"), "dxa")


def set_font(run, size=12, bold=False, italic=False, color="000000"):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(1.25)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    style_specs = {
        "Heading 1": (12, True, False, Pt(18), Pt(18), True),
        "Heading 2": (12, False, False, Pt(18), Pt(6), False),
        "Heading 3": (12, True, False, Pt(12), Pt(6), False),
    }
    for name, (size, bold, italic, before, after, page_break) in style_specs.items():
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.italic = italic
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.left_indent = Cm(0)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_before = before
        style.paragraph_format.space_after = after
        style.paragraph_format.page_break_before = page_break
        style.paragraph_format.keep_with_next = True

    if "Legenda" not in [style.name for style in doc.styles]:
        legend = doc.styles.add_style("Legenda", WD_STYLE_TYPE.PARAGRAPH)
    else:
        legend = doc.styles["Legenda"]
    legend.font.name = "Arial"
    legend.font.size = Pt(10)
    legend.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    legend.paragraph_format.first_line_indent = Cm(0)
    legend.paragraph_format.line_spacing = 1
    legend.paragraph_format.space_before = Pt(3)
    legend.paragraph_format.space_after = Pt(3)

    if "Fonte" not in [style.name for style in doc.styles]:
        source = doc.styles.add_style("Fonte", WD_STYLE_TYPE.PARAGRAPH)
    else:
        source = doc.styles["Fonte"]
    source.font.name = "Arial"
    source.font.size = Pt(10)
    source.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    source.paragraph_format.first_line_indent = Cm(0)
    source.paragraph_format.line_spacing = 1
    source.paragraph_format.space_before = Pt(0)
    source.paragraph_format.space_after = Pt(9)

    reference = doc.styles["Normal"]
    if "Referência ABNT" not in [style.name for style in doc.styles]:
        reference = doc.styles.add_style("Referência ABNT", WD_STYLE_TYPE.PARAGRAPH)
    else:
        reference = doc.styles["Referência ABNT"]
    reference.font.name = "Arial"
    reference.font.size = Pt(12)
    reference.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    reference.paragraph_format.first_line_indent = Cm(0)
    reference.paragraph_format.line_spacing = 1
    reference.paragraph_format.space_after = Pt(12)


def clear_document_body(doc):
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def configure_section(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.header_distance = Cm(2)
    section.footer_distance = Cm(1.25)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_end])
    set_font(run, size=10)


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def clear_header(header):
    for child in list(header._element):
        header._element.remove(child)
    header._element.append(OxmlElement("w:p"))


def set_page_number_start(section, start):
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def add_centered(doc, text="", size=12, bold=False, upper=False, space_after=0):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.space_after = Pt(space_after)
    run = paragraph.add_run(text.upper() if upper else text)
    set_font(run, size=size, bold=bold)
    return paragraph


def add_body(doc, text, bold_lead=None):
    paragraph = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        run = paragraph.add_run(bold_lead)
        set_font(run, bold=True)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_font(rest)
    else:
        run = paragraph.add_run(text)
        set_font(run)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    clean_text = re.sub(r"^\d+(?:\.\d+)*\s+", "", text)
    run = paragraph.add_run(clean_text.upper() if level <= 2 else clean_text)
    set_font(run, bold=(level in (1, 3)))
    return paragraph


def add_unnumbered_heading(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.page_break_before = True
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(18)
    paragraph.paragraph_format.space_after = Pt(18)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text.upper())
    set_font(run, bold=True)
    return paragraph


def add_bullet(doc, text):
    global ALPHA_NUM_ID
    if ALPHA_NUM_ID is None:
        numbering = doc.part.numbering_part.element
        abstract_ids = [
            int(node.get(qn("w:abstractNumId")))
            for node in numbering.findall(qn("w:abstractNum"))
        ]
        num_ids = [
            int(node.get(qn("w:numId")))
            for node in numbering.findall(qn("w:num"))
        ]
        abstract_id = max(abstract_ids, default=0) + 1
        ALPHA_NUM_ID = max(num_ids, default=0) + 1

        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        level.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "lowerLetter")
        level.append(num_fmt)
        level_text = OxmlElement("w:lvlText")
        level_text.set(qn("w:val"), "%1)")
        level.append(level_text)
        level_jc = OxmlElement("w:lvlJc")
        level_jc.set(qn("w:val"), "left")
        level.append(level_jc)
        p_pr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "709")
        ind.set(qn("w:hanging"), "354")
        p_pr.append(ind)
        level.append(p_pr)
        abstract.append(level)
        numbering.append(abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(ALPHA_NUM_ID))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)

    paragraph = doc.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), str(ALPHA_NUM_ID))
    num_pr.extend([ilvl, num_id])
    p_pr.append(num_pr)
    paragraph.paragraph_format.left_indent = Cm(1.25)
    paragraph.paragraph_format.first_line_indent = Cm(-0.63)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = paragraph.add_run(text)
    set_font(run)
    return paragraph


def add_equation(doc, expression):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run(expression)
    set_font(run, bold=True)
    return paragraph


def add_figure(doc, image_path, number, caption, width=15.5):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(image_path), width=Cm(width))
    cap = doc.add_paragraph(style="Legenda")
    run = cap.add_run(f"Figura {number} – {caption}")
    set_font(run, size=10)
    source = doc.add_paragraph(style="Fonte")
    run = source.add_run("Fonte: elaboração dos autores (2026).")
    set_font(run, size=10)


def add_table(doc, headers, rows, widths, caption=None, number=None):
    if caption and number:
        cap = doc.add_paragraph(style="Legenda")
        cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cap.paragraph_format.keep_with_next = True
        run = cap.add_run(f"Tabela {number} – {caption}")
        set_font(run, size=10)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_shading(cell, "D9EAF2")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.line_spacing = 1
            for run in paragraph.runs:
                set_font(run, size=10, bold=True)
    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            cells[index].text = str(value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[index])
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.paragraph_format.line_spacing = 1
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_font(run, size=9.5)
    set_table_widths(table, widths)
    if caption:
        source = doc.add_paragraph(style="Fonte")
        run = source.add_run("Fonte: elaboração dos autores (2026).")
        set_font(run, size=10)
    return table


def create_architecture_diagram():
    path = ASSETS / "arquitetura-software.png"
    canvas = Image.new("RGB", (1500, 570), "white")
    draw = ImageDraw.Draw(canvas)
    try:
        font_bold = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 32)
        font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 25)
    except OSError:
        font_bold = ImageFont.load_default()
        font = ImageFont.load_default()
    boxes = [
        (40, 165, 310, 395, "Interface (HUD)", "Parâmetros\nMedições\nRelatório"),
        (400, 165, 680, 395, "Regras", "Critérios\nCorreções\nDiagnóstico"),
        (770, 165, 1050, 395, "Construtores", "Entrada\nCorredor\nBanheiro"),
        (1140, 165, 1450, 395, "Renderização", "A-Frame\nThree.js\nWebXR"),
    ]
    colors = ["#DCEEF7", "#E5F1E8", "#FFF0D6", "#E8E4F5"]
    for index, (x1, y1, x2, y2, title, body) in enumerate(boxes):
        draw.rounded_rectangle((x1, y1, x2, y2), radius=24, fill=colors[index], outline="#263746", width=4)
        draw.text(((x1 + x2) / 2, y1 + 55), title, fill="#152735", font=font_bold, anchor="mm")
        draw.multiline_text(((x1 + x2) / 2, y1 + 145), body, fill="#263746", font=font,
                            anchor="mm", align="center", spacing=9)
        if index < len(boxes) - 1:
            start_x = x2 + 12
            end_x = boxes[index + 1][0] - 12
            mid_y = (y1 + y2) // 2
            draw.line((start_x, mid_y, end_x, mid_y), fill="#247BA0", width=8)
            draw.polygon([(end_x, mid_y), (end_x - 24, mid_y - 15), (end_x - 24, mid_y + 15)],
                         fill="#247BA0")
    draw.text((750, 60), "Fluxo modular do protótipo", fill="#152735", font=font_bold, anchor="mm")
    draw.text((750, 500), "Alterações de parâmetros reconstruem a cena e recalculam os critérios em tempo real.",
              fill="#263746", font=font, anchor="mm")
    canvas.save(path)
    return path


def create_corridor_comparison():
    path = ASSETS / "corredor-comparacao.png"
    before = Image.open(ASSETS / "corredor-inicial.png").convert("RGB")
    after = Image.open(ASSETS / "corredor-corrigido.png").convert("RGB")
    width = min(before.width, after.width)
    height = min(before.height, after.height)
    before = before.resize((width, height))
    after = after.resize((width, height))
    label_height = 54
    canvas = Image.new("RGB", (width * 2, height + label_height), "white")
    canvas.paste(before, (0, label_height))
    canvas.paste(after, (width, label_height))
    draw = ImageDraw.Draw(canvas)
    try:
        font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 26)
    except OSError:
        font = ImageFont.load_default()
    draw.rectangle((0, 0, width, label_height), fill="#8F3434")
    draw.rectangle((width, 0, width * 2, label_height), fill="#2D7450")
    draw.text((width / 2, label_height / 2), "CENÁRIO INICIAL · 0,25 m",
              fill="white", font=font, anchor="mm")
    draw.text((width * 1.5, label_height / 2), "CENÁRIO CORRIGIDO · 0,10 m",
              fill="white", font=font, anchor="mm")
    canvas.save(path)
    return path


def add_cover(doc):
    add_centered(doc, "UNIVERSIDADE DO SUL DE SANTA CATARINA", bold=True, upper=True)
    add_centered(doc, "UNIDADE CURRICULAR: COMPUTAÇÃO GRÁFICA E REALIDADE VIRTUAL", bold=True, upper=True)
    for _ in range(5):
        doc.add_paragraph()
    add_centered(doc, "[NOMES DOS INTEGRANTES]", bold=True, upper=True)
    for _ in range(5):
        doc.add_paragraph()
    add_centered(doc, TITLE, bold=True, upper=True)
    add_centered(doc, SUBTITLE, upper=False)
    for _ in range(8):
        doc.add_paragraph()
    add_centered(doc, "[CIDADE]", upper=True)
    add_centered(doc, "2026")
    doc.add_page_break()


def add_title_page(doc):
    add_centered(doc, "[NOMES DOS INTEGRANTES]", bold=True, upper=True)
    for _ in range(5):
        doc.add_paragraph()
    add_centered(doc, TITLE, bold=True, upper=True)
    add_centered(doc, SUBTITLE)
    for _ in range(4):
        doc.add_paragraph()
    nature = doc.add_paragraph()
    nature.paragraph_format.left_indent = Cm(8)
    nature.paragraph_format.first_line_indent = Cm(0)
    nature.paragraph_format.line_spacing = 1
    nature.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = nature.add_run(
        "Projeto semestral apresentado à Unidade Curricular de Computação Gráfica e Realidade Virtual "
        "da Universidade do Sul de Santa Catarina, como parte dos requisitos de avaliação do semestre 2026/1."
    )
    set_font(run, size=10)
    nature.add_run("\n\nProfessores: Adalberto Gassenferth Junior e Marcos Tonon.")
    for run in nature.runs:
        set_font(run, size=10)
    for _ in range(6):
        doc.add_paragraph()
    add_centered(doc, "[CIDADE]", upper=True)
    add_centered(doc, "2026")
    doc.add_page_break()


def add_abstract(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Cm(0)
    run = title.add_run("RESUMO")
    set_font(run, bold=True)
    add_body(
        doc,
        "Este trabalho apresenta um protótipo imersivo de pré-auditoria de acessibilidade destinado à "
        "análise preliminar de ambientes construídos antes da execução da obra. A solução representa, em "
        "escala geométrica de uma unidade tridimensional para um metro, três situações recorrentes de "
        "projeto: entrada com rampa, corredor e banheiro acessível. O usuário informa dimensões do "
        "anteprojeto, percorre os ambientes em primeira pessoa e compara os valores modelados com critérios "
        "selecionados de acessibilidade. O sistema calcula inclinação de rampas, larguras livres, vãos de "
        "porta, área de giro e características de equipamentos, além de reconstruir automaticamente uma "
        "alternativa com correções mínimas. A implementação utiliza A-Frame, Three.js, WebGL e WebXR, "
        "modelagem procedural, detecção de colisões e uma interface paramétrica. Nos cenários de teste, a "
        "entrada passou de um para cinco critérios atendidos, o corredor de zero para três e o banheiro "
        "de um para seis. O protótipo demonstra como a computação gráfica e a realidade virtual podem apoiar "
        "a comunicação e a revisão de anteprojetos, sem substituir levantamento técnico, projeto executivo "
        "ou verificação integral da ABNT NBR 9050."
    )
    keywords = doc.add_paragraph()
    keywords.paragraph_format.first_line_indent = Cm(0)
    run = keywords.add_run("Palavras-chave: ")
    set_font(run, bold=True)
    run = keywords.add_run("acessibilidade; realidade virtual; computação gráfica; pré-auditoria; modelagem paramétrica.")
    set_font(run)
    doc.add_page_break()


def add_manual_toc(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Cm(0)
    run = title.add_run("SUMÁRIO")
    set_font(run, bold=True)
    entries = [
        ("1 INTRODUÇÃO", "4"),
        ("1.1 Problema de engenharia e justificativa", "4"),
        ("1.2 Objetivos", "5"),
        ("1.3 Delimitação do trabalho", "5"),
        ("2 PROJETO", "6"),
        ("2.1 Solução proposta e requisitos", "6"),
        ("2.2 Fundamentação em computação gráfica e realidade virtual", "6"),
        ("2.3 Arquitetura e organização do software", "7"),
        ("2.4 Modelagem paramétrica dos ambientes", "7"),
        ("2.5 Critérios e cálculos implementados", "8"),
        ("2.6 Navegação, colisões e interface", "9"),
        ("2.7 Metodologia de desenvolvimento e testes", "10"),
        ("3 RESULTADOS", "11"),
        ("3.1 Entrada e rampa", "11"),
        ("3.2 Corredor", "12"),
        ("3.3 Banheiro", "13"),
        ("3.4 Aplicabilidade e limitações", "14"),
        ("4 CONCLUSÃO", "15"),
        ("REFERÊNCIAS", "16"),
        ("APÊNDICE A – Roteiro sugerido para vídeo de oito minutos", "17"),
    ]
    for label, page in entries:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(15.5))
        run = paragraph.add_run(label)
        set_font(run, bold=label[0].isdigit() and "." not in label.split(" ")[0])
        run = paragraph.add_run("\t" + page)
        set_font(run)


def build():
    doc = Document(TEMPLATE)
    clear_document_body(doc)
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)

    add_cover(doc)
    add_title_page(doc)
    add_abstract(doc)
    add_manual_toc(doc)

    content_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(content_section)
    set_page_number_start(content_section, 4)

    add_heading(doc, "1 INTRODUÇÃO", 1)
    add_body(
        doc,
        "A acessibilidade em edificações é uma condição de uso seguro e autônomo do espaço construído. "
        "Entretanto, incompatibilidades dimensionais ainda podem passar despercebidas durante a leitura de "
        "plantas bidimensionais e ser percebidas somente na obra concluída ou durante uma vistoria. Nessa "
        "fase, ampliar uma porta, reduzir a inclinação de uma rampa ou reorganizar um banheiro pode exigir "
        "demolição, aquisição de novos componentes, alteração de instalações e aumento do prazo de entrega."
    )
    add_body(
        doc,
        "O problema de engenharia abordado é, portanto, a identificação tardia de barreiras de "
        "acessibilidade em anteprojetos. O protótipo desenvolvido transforma parâmetros numéricos em ambientes "
        "tridimensionais navegáveis, permitindo que projetistas, estudantes e usuários observem a relação entre "
        "medidas, circulação e disposição dos elementos antes da execução física."
    )

    add_heading(doc, "1.1 Problema de engenharia e justificativa", 2)
    add_body(
        doc,
        "A representação convencional informa dimensões, mas nem sempre comunica de forma intuitiva a "
        "experiência espacial de uma pessoa com mobilidade reduzida. Uma largura aparentemente próxima do "
        "mínimo pode impedir a circulação quando combinada com saliências, equipamentos ou uma abertura de "
        "porta. A realidade virtual acrescenta a percepção em escala humana, enquanto o cálculo paramétrico "
        "fornece uma verificação objetiva dos valores informados."
    )
    add_body(
        doc,
        "A relevância da solução está na possibilidade de comparar alternativas durante a fase de projeto, "
        "quando alterações são menos dispendiosas. O sistema não emite laudo e não substitui o profissional "
        "responsável, mas funciona como instrumento de pré-auditoria, comunicação e ensino. Essa delimitação "
        "evita transformar uma simulação acadêmica em uma alegação indevida de conformidade normativa."
    )

    add_heading(doc, "1.2 Objetivos", 2)
    add_heading(doc, "1.2.1 Objetivo geral", 3)
    add_body(
        doc,
        "Desenvolver uma aplicação tridimensional imersiva capaz de representar ambientes em escala 1:1, "
        "identificar barreiras dimensionais selecionadas e simular correções de acessibilidade antes da obra."
    )
    add_heading(doc, "1.2.2 Objetivos específicos", 3)
    for item in [
        "modelar proceduralmente uma entrada com rampa, um corredor e um banheiro;",
        "permitir a alteração de medidas reais por meio de uma interface paramétrica;",
        "calcular critérios selecionados de circulação, inclinação, portas e áreas de manobra;",
        "possibilitar navegação em primeira pessoa com colisões e altura de câmera de 1,20 m;",
        "gerar automaticamente uma alternativa corrigida e um diagnóstico textual;",
        "demonstrar a aplicação de WebGL, WebXR, transformações geométricas e interação humano-computador.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "1.3 Delimitação do trabalho", 2)
    add_body(
        doc,
        "O escopo contempla somente os parâmetros disponíveis nos três ambientes do protótipo. Os critérios "
        "são simplificados para fins educacionais e devem ser conferidos na versão oficial da norma e no "
        "contexto específico de cada edificação. O sistema não importa plantas ou modelos BIM, não realiza "
        "levantamento do local, não verifica todos os detalhes construtivos e não considera instalações, "
        "estrutura, incêndio, orçamento ou responsabilidade técnica."
    )

    add_heading(doc, "2 PROJETO", 1)
    add_heading(doc, "2.1 Solução proposta e requisitos", 2)
    add_body(
        doc,
        "A solução foi concebida como uma aplicação web local, executável em navegador compatível com WebGL. "
        "O usuário seleciona um ambiente, informa as dimensões do anteprojeto e observa imediatamente a "
        "reconstrução da cena. A aba de medições apresenta o valor modelado, o critério adotado, o resultado e "
        "uma recomendação. O comando de correção mínima altera os parâmetros necessários e reconstrói a "
        "alternativa para comparação."
    )
    add_table(
        doc,
        ["Requisito", "Implementação", "Finalidade"],
        [
            ("Escala mensurável", "1 unidade 3D = 1 metro", "Relacionar o modelo às dimensões do projeto"),
            ("Interação", "Campos numéricos, chaves e botões", "Testar alternativas sem alterar o código"),
            ("Imersão", "Câmera em primeira pessoa e WebXR", "Avaliar o espaço em perspectiva humana"),
            ("Segurança da navegação", "Colisores e bloqueios de rota", "Simular a mobilidade de uma cadeira de rodas"),
            ("Diagnóstico", "Comparação e recomendações", "Apoiar a revisão preliminar do anteprojeto"),
            ("Usabilidade", "HUD minimizável e redimensionável", "Preservar a visão do ambiente"),
        ],
        [4.0, 5.2, 6.3],
        "Requisitos funcionais do protótipo",
        1,
    )

    add_heading(doc, "2.2 Fundamentação em computação gráfica e realidade virtual", 2)
    add_body(
        doc,
        "A computação gráfica é aplicada na criação de primitivas, materiais, iluminação, câmera e malhas "
        "tridimensionais. Cada ambiente é construído a partir de caixas, cilindros, planos, esferas e uma malha "
        "triangular específica para a rampa. As posições são expressas por coordenadas cartesianas nos eixos x, "
        "y e z; dimensões e transformações são recalculadas quando um parâmetro é modificado."
    )
    add_body(
        doc,
        "O A-Frame fornece a estrutura de entidades e componentes e integra a cena à WebXR. O Three.js é "
        "utilizado para vetores, quaternions, geometrias e renderização WebGL. Para orientar uma barra entre "
        "dois pontos, o programa calcula o vetor direção, o ponto médio, o comprimento e a rotação que alinha "
        "o eixo local do cilindro ao segmento. Esse procedimento é empregado nos corrimãos, barras de apoio, "
        "travessas e indicadores."
    )
    add_equation(doc, "ponto médio = (ponto inicial + ponto final) / 2")
    add_equation(doc, "comprimento = ‖ponto final − ponto inicial‖")
    add_body(
        doc,
        "A realidade virtual não é utilizada apenas como acabamento visual. Ela permite analisar aproximação, "
        "campo de visão, circulação e relação entre os elementos na escala do usuário. A câmera foi posicionada "
        "a 1,20 m para representar aproximadamente a perspectiva de uma pessoa sentada."
    )

    add_heading(doc, "2.3 Arquitetura e organização do software", 2)
    add_body(
        doc,
        "O código foi dividido por responsabilidade. O núcleo geométrico centraliza a criação de primitivas e "
        "segmentos orientados; cada ambiente possui seu próprio construtor; o catálogo de ambientes concentra "
        "parâmetros, regras e correções; a navegação processa teclado, controle analógico e colisões; e o módulo "
        "de interface mantém o estado, atualiza o diagnóstico e exporta o relatório."
    )
    diagram = create_architecture_diagram()
    add_figure(doc, diagram, 1, "Arquitetura modular da aplicação")
    add_body(
        doc,
        "Essa separação reduz o acoplamento e permite modificar um ambiente sem duplicar o núcleo geométrico. "
        "Também evita que a orientação de barras, anteriormente repetida em diferentes trechos, seja calculada "
        "de formas incompatíveis."
    )

    add_heading(doc, "2.4 Modelagem paramétrica dos ambientes", 2)
    add_heading(doc, "2.4.1 Entrada e rampa", 3)
    add_body(
        doc,
        "A entrada recebe desnível, comprimento e largura da rampa, vão livre da porta, ressalto e presença de "
        "corrimãos. A rampa é uma malha em forma de prisma triangular, e sua superfície é usada também para "
        "calcular a altura da câmera durante o deslocamento. Rampa e patamar compartilham a mesma cota no ponto "
        "de encontro, evitando degrau geométrico. Os corrimãos são construídos bilateralmente em duas alturas, "
        "conectados por montantes e travessas, e possuem barreiras de colisão contínuas."
    )
    add_heading(doc, "2.4.2 Corredor", 3)
    add_body(
        doc,
        "O corredor permite alterar a largura livre, a profundidade de uma saliência e a existência de "
        "sinalização tátil. O abrigo do extintor e o quadro de serviço permanecem modelados em todos os "
        "cenários. Quando reduzidos a 0,10 m, passam a ocupar uma faixa técnica junto à parede e deixam livre "
        "a circulação; acima desse valor, continuam visíveis, são destacados como inadequados e recebem "
        "volumes de colisão. Assim, conformidade não é representada pela remoção artificial do equipamento."
    )
    add_heading(doc, "2.4.3 Banheiro", 3)
    add_body(
        doc,
        "O banheiro combina largura, profundidade, diâmetro de giro, vão da porta, altura e tipo de lavatório e "
        "barras de apoio. A porta abre para fora, a bacia está orientada para a área de uso e o lavatório possui "
        "volume de colisão; a área de manobra é representada no piso. O arranjo funcional adotado no protótipo "
        "exige 2,60 m por 2,50 m para "
        "evitar sobreposição entre o círculo de giro e as peças desta configuração específica; esse valor é uma "
        "decisão de layout do modelo, não um mínimo universal atribuído à norma."
    )

    add_heading(doc, "2.5 Critérios e cálculos implementados", 2)
    add_heading(doc, "2.5.1 Inclinação e comprimento da rampa", 3)
    add_body(
        doc,
        "A inclinação percentual é obtida pela razão entre o desnível vertical h e o comprimento horizontal L. "
        "No cenário inicial, h = 0,50 m e L = 4,50 m."
    )
    add_equation(doc, "inclinação = (h / L) × 100 = (0,50 / 4,50) × 100 = 11,11%")
    add_body(
        doc,
        "Considerando no protótipo o limite selecionado de 8,33%, o comprimento mínimo é calculado por:"
    )
    add_equation(doc, "L mínimo = h / 0,0833 = 0,50 / 0,0833 ≈ 6,01 m")
    add_body(
        doc,
        "Após a correção, o comprimento passa a 6,01 m e a inclinação calculada é aproximadamente 8,32%."
    )

    add_heading(doc, "2.5.2 Circulação, portas e manobra", 3)
    doc.paragraphs[-1].paragraph_format.page_break_before = True
    add_table(
        doc,
        ["Item", "Cenário inicial", "Critério adotado", "Correção"],
        [
            ("Largura da rampa", "1,00 m", "≥ 1,20 m", "1,20 m"),
            ("Vão da porta da entrada", "0,65 m", "≥ 0,80 m", "0,80 m"),
            ("Ressalto", "0,00 m", "≤ 0,005 m", "0,00 m"),
            ("Largura do corredor", "0,90 m", "≥ 1,20 m", "1,20 m"),
            ("Saliência no corredor", "0,25 m", "≤ 0,10 m", "0,10 m"),
            ("Giro no banheiro", "Ø 1,10 m", "Ø ≥ 1,50 m", "Ø 1,50 m"),
            ("Vão da porta do banheiro", "0,70 m", "≥ 0,80 m", "0,80 m"),
            ("Altura do lavatório", "0,95 m", "0,80–0,85 m", "0,83 m"),
        ],
        [4.2, 3.5, 3.8, 3.8],
        "Parâmetros numéricos utilizados nos testes",
        2,
    )
    add_body(
        doc,
        "As condições booleanas verificam a presença de corrimãos, barras de apoio, sinalização tátil e espaço "
        "livre sob o lavatório. A correção automática não procura uma solução ótima de projeto; ela aplica os "
        "valores mínimos configurados no protótipo para produzir uma alternativa comparável."
    )

    add_heading(doc, "2.6 Navegação, colisões e interface", 2)
    add_body(
        doc,
        "O deslocamento utiliza as teclas W, A, S e D, setas direcionais ou os eixos de um controle. A direção "
        "é calculada a partir da orientação da câmera e limitada ao plano horizontal. Antes de atualizar a "
        "posição, o sistema testa separadamente os eixos x e z contra limites, caixas e círculos de colisão. "
        "Cada teste retorna um resultado com permissão e motivo do bloqueio, permitindo deslizar junto a "
        "paredes sem atravessá-las e informar ao usuário a barreira encontrada."
    )
    add_body(
        doc,
        "A navegação representa o perfil de uma pessoa em cadeira de rodas. Se inclinação, largura, porta, "
        "soleira ou corrimãos não atenderem às regras configuradas, a base da rampa é bloqueada. As laterais "
        "possuem colisores contínuos e a face elevada do patamar impede a subida por fora da rampa. Após a "
        "correção, a faixa central torna-se transitável e a altura da câmera varia continuamente entre o piso, "
        "a rampa e o patamar."
    )
    add_body(
        doc,
        "O HUD apresenta parâmetros, medições e informações do projeto. Ele pode ser minimizado e redimensionado "
        "por arraste ou teclado, com persistência local do tamanho. A remoção de textos flutuantes da cena evita "
        "oclusão do ambiente. No banheiro, a recentralização posiciona a câmera dentro da faixa livre para que "
        "a inspeção não seja bloqueada pela folha da porta ou pelo lavatório."
    )

    add_heading(doc, "2.7 Metodologia de desenvolvimento e testes", 2)
    add_body(
        doc,
        "O desenvolvimento foi incremental. Primeiro foram construídos os ambientes e a navegação; em seguida, "
        "foram introduzidos escala, parâmetros, colisões, diagnóstico e correção automática. As revisões visuais "
        "identificaram problemas como um elemento vertical indevido no centro da rampa. A causa foi rastreada ao "
        "ciclo de atualização do A-Frame, que sobrescrevia rotações aplicadas diretamente ao objeto Three.js. A "
        "solução foi centralizar a criação de segmentos no núcleo geométrico e definir posição e rotação como "
        "atributos da entidade."
    )
    add_body(
        doc,
        "Foram realizadas checagens de sintaxe em todos os módulos JavaScript, inspeções no navegador e testes "
        "automatizados com a biblioteca nativa de testes do Node.js. Os casos cobrem o bloqueio da rampa não "
        "conforme, a colisão dos corrimãos, a impossibilidade de escalar a face lateral do patamar, a "
        "continuidade da rampa corrigida e a permanência dos equipamentos do corredor em 0,10 m. Os quatro "
        "testes passaram, e os cenários corrigidos atingiram 5 de 5 critérios na entrada, 3 de 3 no corredor e "
        "6 de 6 no banheiro. Esses resultados confirmam a coerência interna das regras, mas não equivalem a "
        "validação externa ou certificação normativa."
    )

    add_heading(doc, "3 RESULTADOS", 1)
    add_heading(doc, "3.1 Entrada e rampa", 2)
    add_body(
        doc,
        "No cenário inicial, a rampa apresentou inclinação de 11,11%, largura de 1,00 m, porta de 0,65 m, "
        "soleira nivelada e ausência de corrimãos. Apenas o critério de ressalto foi atendido, resultando em "
        "um de cinco critérios. A rota permaneceu bloqueada para o perfil cadeirante, pois inclinação, largura, "
        "porta e proteção lateral ainda eram inadequadas."
    )
    add_figure(doc, ASSETS / "entrada-inicial.png", 2, "Entrada no cenário inicial")
    add_body(
        doc,
        "A correção automática ampliou o comprimento para 6,01 m, a largura para 1,20 m e a porta para 0,80 m, "
        "manteve a soleira nivelada e inseriu corrimãos bilaterais conectados em duas alturas. Também alinhou "
        "geometricamente o topo da rampa ao patamar e habilitou a rota central. O resultado passou a cinco "
        "critérios atendidos."
    )
    add_figure(doc, ASSETS / "entrada-corrigida.png", 3, "Entrada após a aplicação das correções")

    add_heading(doc, "3.2 Corredor", 2)
    add_body(
        doc,
        "O corredor inicial possuía 0,90 m de largura, saliência de 0,25 m e ausência de sinalização tátil, "
        "resultando em zero de três critérios atendidos. A alternativa corrigida ampliou a faixa para 1,20 m, "
        "reduziu a projeção dos equipamentos para 0,10 m e adicionou uma faixa tátil contrastante. O abrigo do "
        "extintor e o quadro de serviço não desaparecem: permanecem visíveis, parcialmente embutidos na faixa "
        "técnica da parede e fora da circulação útil. Acima do limite, cada equipamento possui colisor próprio."
    )
    corridor_comparison = create_corridor_comparison()
    add_figure(doc, corridor_comparison, 4, "Equipamentos do corredor antes e depois da correção", width=15.5)
    add_table(
        doc,
        ["Indicador", "Antes", "Depois"],
        [
            ("Critérios atendidos – entrada", "1/5", "5/5"),
            ("Critérios atendidos – corredor", "0/3", "3/3"),
            ("Critérios atendidos – banheiro", "1/6", "6/6"),
            ("Diagnóstico", "Não conforme ou parcial", "Conforme às regras implementadas"),
        ],
        [7.5, 4.0, 4.0],
        "Síntese dos cenários testados",
        3,
    )

    add_heading(doc, "3.3 Banheiro", 2)
    add_body(
        doc,
        "O banheiro inicial atendia apenas às dimensões gerais simplificadas, mas falhava no giro, no vão da "
        "porta, no lavatório, nas barras e no arranjo funcional da planta. O círculo vermelho evidenciou a área "
        "insuficiente, e a navegação permitiu observar a proximidade entre peças."
    )
    add_figure(doc, ASSETS / "banheiro-inicial.png", 5, "Banheiro no cenário inicial")
    add_body(
        doc,
        "A correção ampliou o ambiente para 2,60 m por 2,50 m, adotou giro de 1,50 m, porta de 0,80 m, lavatório "
        "suspenso a 0,83 m e barras horizontais conectadas nas paredes lateral e de fundo. O usuário passa a "
        "iniciar dentro do ambiente para percorrer a área e inspecionar a configuração."
    )
    add_figure(doc, ASSETS / "banheiro-corrigido.png", 6, "Banheiro após a aplicação das correções")

    add_heading(doc, "3.4 Aplicabilidade e limitações", 2)
    add_body(
        doc,
        "A principal aplicação é a revisão preliminar e a comunicação de anteprojetos. Um projetista pode "
        "apresentar ao cliente ou à equipe uma condição inicial, modificar dimensões e demonstrar o efeito da "
        "correção. Em contexto educacional, o sistema relaciona fórmulas e critérios com uma experiência "
        "espacial. Em uma evolução profissional, os parâmetros poderiam ser obtidos de modelos IFC ou glTF e "
        "as verificações poderiam registrar o item normativo, a versão da regra, evidências e responsáveis."
    )
    add_body(
        doc,
        "As limitações incluem o número reduzido de ambientes, simplificação visual, ausência de levantamento "
        "automático, conjunto parcial de critérios, aproximações geométricas nos colisores e falta de validação com "
        "profissionais e usuários reais. Além disso, a norma de acessibilidade contém condições e exceções que "
        "não podem ser reduzidas a um único valor sem análise do contexto. Por isso, a indicação “conforme” da "
        "interface significa somente conformidade com as regras implementadas naquele cenário."
    )

    add_heading(doc, "4 CONCLUSÃO", 1)
    add_body(
        doc,
        "O trabalho demonstrou que técnicas de computação gráfica e realidade virtual podem ser aplicadas a um "
        "problema real de engenharia: a identificação tardia de barreiras de acessibilidade. A combinação de "
        "modelagem procedural, parâmetros em metros, cálculos, colisões e navegação imersiva produziu uma "
        "ferramenta capaz de comparar alternativas antes da execução física."
    )
    add_body(
        doc,
        "Os objetivos propostos foram alcançados dentro do escopo acadêmico. Os três ambientes respondem às "
        "alterações do usuário, apresentam diagnóstico, aplicam correções e permitem inspeção em primeira "
        "pessoa. Os testes internos mostraram a passagem de cenários com falhas para cenários que atendem a "
        "todas as regras programadas."
    )
    add_body(
        doc,
        "A contribuição não está em substituir o projeto ou o laudo, mas em antecipar discussões, tornar "
        "dimensões compreensíveis e apoiar decisões preliminares. Como trabalhos futuros, recomenda-se importar "
        "modelos BIM, ampliar a cobertura normativa, expandir os testes automatizados de geometria, registrar "
        "rastreabilidade dos critérios e realizar avaliações com engenheiros, arquitetos e pessoas com "
        "deficiência ou mobilidade reduzida."
    )

    add_unnumbered_heading(doc, "REFERÊNCIAS")
    references = [
        "ASSOCIAÇÃO BRASILEIRA DE NORMAS TÉCNICAS. ABNT NBR 9050: acessibilidade a edificações, mobiliário, espaços e equipamentos urbanos. Rio de Janeiro: ABNT, 2020.",
        "A-FRAME. A-Frame documentation. [S. l.], 2026. Disponível em: https://aframe.io/docs/. Acesso em: 14 jun. 2026.",
        "MDN WEB DOCS. WebXR Device API. [S. l.], 2026. Disponível em: https://developer.mozilla.org/en-US/docs/Web/API/WebXR_Device_API. Acesso em: 14 jun. 2026.",
        "THREE.JS. Three.js documentation. [S. l.], 2026. Disponível em: https://threejs.org/docs/. Acesso em: 14 jun. 2026.",
        "W3C. WebXR Device API. [S. l.], 2026. Disponível em: https://www.w3.org/TR/webxr/. Acesso em: 14 jun. 2026.",
    ]
    for reference in references:
        paragraph = doc.add_paragraph(style="Referência ABNT")
        run = paragraph.add_run(reference)
        set_font(run)

    add_unnumbered_heading(doc, "APÊNDICE A – ROTEIRO SUGERIDO PARA VÍDEO DE OITO MINUTOS")
    add_table(
        doc,
        ["Tempo", "Conteúdo", "Demonstração sugerida"],
        [
            ("0:00–0:50", "Problema de engenharia", "Retrabalho e dificuldade de perceber barreiras em plantas 2D"),
            ("0:50–1:30", "Solução proposta", "Apresentar a pré-auditoria paramétrica e imersiva"),
            ("1:30–2:30", "Embasamento", "Escala, modelagem procedural, WebGL, WebXR e colisões"),
            ("2:30–5:40", "Programa", "Entrada, corredor e banheiro antes/depois"),
            ("5:40–6:30", "Programação", "Arquitetura modular e reconstrução em tempo real"),
            ("6:30–7:20", "Resultados", "Mostrar as pontuações 1/5→5/5, 0/3→3/3 e 1/6→6/6"),
            ("7:20–8:00", "Conclusão", "Aplicações, limitações e evolução para BIM"),
        ],
        [2.5, 5.2, 8.0],
        "Distribuição recomendada do tempo de apresentação",
        4,
    )
    add_body(
        doc,
        "Durante a gravação, recomenda-se alternar a fala dos integrantes, manter o HUD visível apenas durante "
        "a configuração e as medições e minimizá-lo durante a navegação. A apresentação deve enfatizar que o "
        "resultado é uma pré-auditoria educacional e que a avaliação profissional exige análise completa da "
        "norma e do contexto da edificação."
    )

    for section in doc.sections:
        configure_section(section)

    sections = doc.sections
    for section in sections:
        section.header.is_linked_to_previous = False
        clear_header(section.header)
    final_header = sections[-1].header
    final_header_paragraph = final_header.paragraphs[0]
    final_header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    final_header_paragraph.paragraph_format.first_line_indent = Cm(0)
    add_page_field(final_header_paragraph)

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    doc.core_properties.title = TITLE
    doc.core_properties.subject = "Projeto semestral de Computação Gráfica e Realidade Virtual"
    doc.core_properties.author = "[NOMES DOS INTEGRANTES]"
    doc.core_properties.keywords = "acessibilidade, realidade virtual, computação gráfica, NBR 9050"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
